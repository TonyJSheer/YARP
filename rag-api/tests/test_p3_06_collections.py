"""Tests for P3-06: Collections, file formats, metadata, and reindex."""

import io
import uuid
from pathlib import Path
from unittest.mock import patch

import pytest
from fastapi.testclient import TestClient

from app.db import SessionLocal
from app.models.chunk import Chunk
from app.models.document import Document
from app.services import ingestion

FAKE_VECTOR = [0.1] * 768
FAKE_VECTOR_NEW = [0.9] * 768


@pytest.fixture(autouse=True)
def mock_embed_chunks() -> pytest.FixtureRequest:
    with patch(
        "app.services.embedding.embed_chunks",
        side_effect=lambda texts: [FAKE_VECTOR for _ in texts],
    ):
        yield


# ---------------------------------------------------------------------------
# Part 1: Collections
# ---------------------------------------------------------------------------


def test_upload_with_collection(
    client: TestClient,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    auth_headers: dict[str, str],
) -> None:
    """Upload with collection='legal', assert collection stored."""
    from app.config import settings

    monkeypatch.setattr(settings, "upload_dir", str(tmp_path))

    response = client.post(
        "/documents",
        data={"collection": "legal"},
        files={"file": ("doc.txt", b"Legal document content.", "text/plain")},
        headers=auth_headers,
    )
    assert response.status_code == 201
    doc_id = uuid.UUID(response.json()["document_id"])

    with SessionLocal() as db:
        doc = db.get(Document, doc_id)
        assert doc is not None
        assert doc.collection == "legal"


def test_upload_default_collection(
    client: TestClient,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    auth_headers: dict[str, str],
) -> None:
    """Upload without collection param — should default to 'default'."""
    from app.config import settings

    monkeypatch.setattr(settings, "upload_dir", str(tmp_path))

    response = client.post(
        "/documents",
        files={"file": ("doc.txt", b"Some content.", "text/plain")},
        headers=auth_headers,
    )
    assert response.status_code == 201
    doc_id = uuid.UUID(response.json()["document_id"])

    with SessionLocal() as db:
        doc = db.get(Document, doc_id)
        assert doc is not None
        assert doc.collection == "default"


def test_list_documents_includes_collection(
    client: TestClient,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    auth_headers: dict[str, str],
) -> None:
    """GET /documents returns collection field on each item."""
    from app.config import settings

    monkeypatch.setattr(settings, "upload_dir", str(tmp_path))

    client.post(
        "/documents",
        data={"collection": "specs"},
        files={"file": ("spec.txt", b"A spec document.", "text/plain")},
        headers=auth_headers,
    )

    resp = client.get("/documents", headers=auth_headers)
    assert resp.status_code == 200
    docs = resp.json()["documents"]
    assert len(docs) >= 1
    # The most recently uploaded doc should have collection='specs'
    assert any(d["collection"] == "specs" for d in docs)


def test_query_scoped_to_collection(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """Chunks from collection 'A' are not returned when querying collection 'B'."""
    from app.config import settings
    from app.services import retrieval

    monkeypatch.setattr(settings, "upload_dir", str(tmp_path))

    account_id = f"coll-scope-{uuid.uuid4().hex[:8]}"

    with SessionLocal() as db:
        doc_a_id, _ = ingestion.ingest_from_bytes(
            "a.txt",
            b"Alpha document content for collection A.",
            account_id,
            db,
            collection="A",
        )
        doc_b_id, _ = ingestion.ingest_from_bytes(
            "b.txt",
            b"Beta document content for collection B.",
            account_id,
            db,
            collection="B",
        )

    # Query scoped to collection "A" — should only return A's chunks
    with SessionLocal() as db:
        results = retrieval.retrieve(
            query_embedding=FAKE_VECTOR,
            account_id=account_id,
            db=db,
            top_k=10,
            search_mode="vector",
            collection="A",
        )

    assert len(results) > 0
    for chunk in results:
        assert chunk.document_id == doc_a_id
        assert chunk.document_id != doc_b_id


def test_list_collections_rest(
    client: TestClient,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    auth_headers: dict[str, str],
) -> None:
    """GET /collections returns collection names and counts."""
    from app.config import settings

    monkeypatch.setattr(settings, "upload_dir", str(tmp_path))

    # Use a unique account to avoid cross-test contamination
    import jwt

    unique_account = f"coll-test-{uuid.uuid4().hex[:8]}"
    token = jwt.encode({"sub": unique_account}, "test-secret", algorithm="HS256")
    headers = {"Authorization": f"Bearer {token}"}

    client.post(
        "/documents",
        data={"collection": "alpha"},
        files={"file": ("a.txt", b"Alpha content.", "text/plain")},
        headers=headers,
    )
    client.post(
        "/documents",
        data={"collection": "beta"},
        files={"file": ("b.txt", b"Beta content.", "text/plain")},
        headers=headers,
    )
    client.post(
        "/documents",
        data={"collection": "alpha"},
        files={"file": ("a2.txt", b"More alpha content.", "text/plain")},
        headers=headers,
    )

    resp = client.get("/collections", headers=headers)
    assert resp.status_code == 200
    body = resp.json()
    assert "collections" in body

    coll_map = {c["name"]: c["document_count"] for c in body["collections"]}
    assert coll_map.get("alpha") == 2
    assert coll_map.get("beta") == 1


# ---------------------------------------------------------------------------
# Part 2: Additional file formats
# ---------------------------------------------------------------------------


def test_docx_extraction(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """Upload a .docx file — chunks are extracted."""
    from app.config import settings

    monkeypatch.setattr(settings, "upload_dir", str(tmp_path))

    # Build a minimal .docx in memory using python-docx
    from docx import Document as DocxDocument

    docx_doc = DocxDocument()
    docx_doc.add_paragraph("This is a test paragraph in a Word document.")
    docx_doc.add_paragraph("Second paragraph with more content for chunking.")
    buf = io.BytesIO()
    docx_doc.save(buf)
    docx_bytes = buf.getvalue()

    account_id = f"docx-{uuid.uuid4().hex[:8]}"
    with SessionLocal() as db:
        doc_id, chunk_count = ingestion.ingest_from_bytes("test.docx", docx_bytes, account_id, db)

    assert chunk_count >= 1
    with SessionLocal() as db:
        chunks = db.query(Chunk).filter(Chunk.document_id == doc_id).all()
        assert len(chunks) >= 1
        assert all("paragraph" in c.text.lower() or "content" in c.text.lower() for c in chunks)


def test_html_extraction(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """Upload an .html file — HTML tags stripped, text extracted."""
    from app.config import settings

    monkeypatch.setattr(settings, "upload_dir", str(tmp_path))

    html_content = b"""<html>
<head><title>Test Page</title></head>
<body>
<nav><a href="/">Home</a></nav>
<h1>Main Article</h1>
<p>This is the article content we want to index.</p>
<footer>Footer text we do not want.</footer>
<script>var x = 1;</script>
</body>
</html>"""

    account_id = f"html-{uuid.uuid4().hex[:8]}"
    with SessionLocal() as db:
        doc_id, chunk_count = ingestion.ingest_from_bytes("page.html", html_content, account_id, db)

    assert chunk_count >= 1
    with SessionLocal() as db:
        chunks = db.query(Chunk).filter(Chunk.document_id == doc_id).all()
        full_text = " ".join(c.text for c in chunks)
        assert "Main Article" in full_text or "article content" in full_text
        assert "<html>" not in full_text
        assert "<p>" not in full_text
        # nav/footer/script stripped
        assert "var x = 1" not in full_text


def test_csv_extraction(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """Upload a .csv file — rows extracted as chunks."""
    from app.config import settings

    monkeypatch.setattr(settings, "upload_dir", str(tmp_path))

    # Build a CSV with 25 rows to get multiple batches
    lines = ["name,value,category"]
    for i in range(25):
        lines.append(f"item_{i},{i * 10},cat_{i % 3}")
    csv_bytes = "\n".join(lines).encode("utf-8")

    account_id = f"csv-{uuid.uuid4().hex[:8]}"
    with SessionLocal() as db:
        doc_id, chunk_count = ingestion.ingest_from_bytes("data.csv", csv_bytes, account_id, db)

    assert chunk_count >= 1
    with SessionLocal() as db:
        chunks = db.query(Chunk).filter(Chunk.document_id == doc_id).all()
        assert len(chunks) >= 1
        full_text = " ".join(c.text for c in chunks)
        assert "item_0" in full_text
        assert "name:" in full_text


# ---------------------------------------------------------------------------
# Part 3: Document metadata
# ---------------------------------------------------------------------------


def test_upload_with_metadata(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """Upload via ingest_from_bytes with metadata — metadata stored and returned."""
    from app.config import settings
    from app.services import document_service

    monkeypatch.setattr(settings, "upload_dir", str(tmp_path))

    account_id = f"meta-{uuid.uuid4().hex[:8]}"
    meta = {"source": "Q4-report", "author": "alice"}

    with SessionLocal() as db:
        doc_id, _ = ingestion.ingest_from_bytes(
            "report.txt",
            b"Quarterly report content.",
            account_id,
            db,
            metadata=meta,
        )

    with SessionLocal() as db:
        items = document_service.list_documents(account_id, db)

    assert len(items) == 1
    assert items[0].metadata == meta


def test_metadata_optional(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """Upload without metadata — no error, metadata is None."""
    from app.config import settings
    from app.services import document_service

    monkeypatch.setattr(settings, "upload_dir", str(tmp_path))

    account_id = f"no-meta-{uuid.uuid4().hex[:8]}"

    with SessionLocal() as db:
        ingestion.ingest_from_bytes("plain.txt", b"Plain document.", account_id, db)

    with SessionLocal() as db:
        items = document_service.list_documents(account_id, db)

    assert len(items) == 1
    assert items[0].metadata is None


# ---------------------------------------------------------------------------
# Part 4: Re-index
# ---------------------------------------------------------------------------


def test_reindex_updates_embeddings(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """reindex_document re-embeds chunks with the current model."""
    from app.config import settings
    from app.services import document_service

    monkeypatch.setattr(settings, "upload_dir", str(tmp_path))

    account_id = f"reindex-{uuid.uuid4().hex[:8]}"

    with SessionLocal() as db:
        doc_id, _ = ingestion.ingest_from_bytes(
            "reindex.txt",
            b"Content to be re-indexed after model change.",
            account_id,
            db,
        )

    # Verify initial embeddings are FAKE_VECTOR
    with SessionLocal() as db:
        chunks_before = db.query(Chunk).filter(Chunk.document_id == doc_id).all()
        assert all(list(c.embedding) == FAKE_VECTOR for c in chunks_before)

    # Re-index with new vectors
    with patch(
        "app.services.embedding.embed_chunks",
        side_effect=lambda texts: [FAKE_VECTOR_NEW for _ in texts],
    ):
        with SessionLocal() as db:
            chunk_count = document_service.reindex_document(str(doc_id), account_id, db)

    assert chunk_count >= 1

    with SessionLocal() as db:
        chunks_after = db.query(Chunk).filter(Chunk.document_id == doc_id).all()
        assert all(list(c.embedding) == FAKE_VECTOR_NEW for c in chunks_after)


def test_reindex_wrong_account_raises_not_found(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """reindex_document raises DocumentNotFoundError for wrong account."""
    from app.config import settings
    from app.services.document_service import DocumentNotFoundError

    monkeypatch.setattr(settings, "upload_dir", str(tmp_path))

    owner_account = f"owner-{uuid.uuid4().hex[:8]}"
    other_account = f"other-{uuid.uuid4().hex[:8]}"

    with SessionLocal() as db:
        doc_id, _ = ingestion.ingest_from_bytes(
            "private.txt", b"Private document.", owner_account, db
        )

    from app.services import document_service

    with SessionLocal() as db:
        with pytest.raises(DocumentNotFoundError):
            document_service.reindex_document(str(doc_id), other_account, db)
